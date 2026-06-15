import json
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt

from app.core.config import get_settings
from app.services.llm.prompts import PROJECT_DOC_SECTIONS, SECTION_LABELS


def build_project_doc(
    *,
    title: str,
    section_data: dict,
    user_email: str,
    session_id: int,
    model_used: str,
) -> Document:
    settings = get_settings()
    doc = Document()

    doc.add_heading(title or "Projeto de Ciência de Dados em Saúde", level=0)
    doc.add_paragraph(f"Instituição: {settings.institution_name}")
    doc.add_paragraph(f"Gerado em: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph(f"Autor: {user_email}")
    doc.add_paragraph(f"ID da Sessão: {session_id}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Finalidade: Documentar os dados, métodos, artefatos esperados e fluxo de análise do estudo."
    )
    doc.add_paragraph("")

    for section_key in PROJECT_DOC_SECTIONS:
        label = SECTION_LABELS.get(section_key, section_key.replace("_", " ").title())
        doc.add_heading(label, level=1)
        content = section_data.get(section_key, "")
        if isinstance(content, dict):
            content = content.get("content", json.dumps(content, indent=2))
        doc.add_paragraph(str(content) if content else "[Não informado]")

    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = (
        f"Exportação HRA | {user_email} | Sessão {session_id} | Modelo: {model_used} | "
        f"{datetime.now(timezone.utc).isoformat()}"
    )
    for run in footer_para.runs:
        run.font.size = Pt(8)

    return doc
